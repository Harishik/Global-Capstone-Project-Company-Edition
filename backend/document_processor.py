"""
Intellecta RAG Backend - Document Processor
Document parsing for PDF, DOCX, CSV, Excel, TXT, MD, JSON with 512-token chunking.
Optimized for energy sector documents with domain-specific preprocessing.
"""

import hashlib
import json
import logging
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import tiktoken

from models import (
    DocumentChunk, ChunkMetadata, DocumentRecord, DocumentStatus,
    SecurityLevel, IngestResponse
)
from security import security_checker

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ===== Configuration =====

CHUNK_SIZE = 512  # tokens
CHUNK_OVERLAP = 50  # tokens
MIN_CHUNK_SIZE = 100  # tokens

UPLOAD_DIR = Path(__file__).parent / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".csv": "csv",
    ".xlsx": "excel",
    ".xls": "excel",
    ".txt": "text",
    ".md": "markdown",
    ".json": "json",
    ".hwp": "hwp",
}

# Energy domain keywords for classification
ENERGY_DOMAINS = {
    "solar": ["solar", "photovoltaic", "pv", "panel", "irradiance", "inverter"],
    "wind": ["wind", "turbine", "rotor", "nacelle", "offshore", "onshore"],
    "grid": ["grid", "transmission", "distribution", "substation", "transformer", "voltage"],
    "power_plant": ["power plant", "generation", "capacity", "megawatt", "mw", "gw"],
    "nuclear": ["nuclear", "reactor", "uranium", "fission"],
    "hydro": ["hydro", "dam", "reservoir", "hydroelectric"],
    "storage": ["battery", "storage", "lithium", "energy storage"],
    "anomaly": ["anomaly", "detection", "outlier", "abnormal", "fault"],
}

# Initialize tokenizer
try:
    tokenizer = tiktoken.get_encoding("cl100k_base")
except:
    tokenizer = None
    logger.warning("tiktoken not available, using approximate token counting")


class DocumentProcessor:
    """Process documents into chunks for vector storage."""
    
    def __init__(self):
        self.chunk_size = CHUNK_SIZE
        self.chunk_overlap = CHUNK_OVERLAP
        self.min_chunk_size = MIN_CHUNK_SIZE
    
    def process_file(
        self,
        file_path: str,
        filename: str,
        security_level: SecurityLevel = SecurityLevel.PUBLIC,
        source: Optional[str] = None
    ) -> Tuple[List[DocumentChunk], str]:
        """
        Process a file and return chunks.
        
        Args:
            file_path: Path to the file
            filename: Original filename
            security_level: Security classification
            source: Dataset source (opsd, nrel, etc.)
            
        Returns:
            Tuple of (chunks, doc_id)
        """
        # Generate document ID
        doc_id = self._generate_doc_id(filename)
        
        # Determine file type
        ext = Path(filename).suffix.lower()
        file_type = SUPPORTED_EXTENSIONS.get(ext)
        
        if not file_type:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Extract text based on file type
        logger.info(f"Processing {filename} as {file_type}")
        
        if file_type == "pdf":
            text = self._extract_pdf(file_path)
        elif file_type == "docx":
            text = self._extract_docx(file_path)
        elif file_type in ["csv", "excel"]:
            text = self._extract_tabular(file_path, file_type)
        elif file_type == "text":
            text = self._extract_text(file_path)
        elif file_type == "markdown":
            text = self._extract_markdown(file_path)
        elif file_type == "json":
            text = self._extract_json(file_path)
        else:
            text = self._extract_text(file_path)
        
        if not text or not text.strip():
            raise ValueError(f"No text content extracted from {filename}")
        
        # Preprocess text (energy domain optimization)
        text = self._preprocess_text(text)
        
        # Detect domain
        domain = self._detect_domain(text)
        
        # Auto-detect security if not specified
        if security_level == SecurityLevel.PUBLIC:
            detected = security_checker.auto_detect_security(text)
            if detected.level_value > 0:
                security_level = detected.detected_level
                logger.info(f"Auto-detected security level: {security_level.value}")
        
        # Chunk the text
        chunks = self._chunk_text(
            text=text,
            doc_id=doc_id,
            filename=filename,
            source=source,
            security_level=security_level,
            domain=domain,
            file_type=file_type
        )
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        
        return chunks, doc_id
    
    def _generate_doc_id(self, filename: str) -> str:
        """Generate unique document ID."""
        timestamp = datetime.now().isoformat()
        unique_str = f"{filename}_{timestamp}_{uuid.uuid4().hex[:8]}"
        return hashlib.md5(unique_str.encode()).hexdigest()[:16]
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            doc.close()
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
            raise
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def _extract_tabular(self, file_path: str, file_type: str) -> str:
        """Extract text from CSV or Excel files."""
        try:
            if file_type == "csv":
                df = pd.read_csv(file_path, nrows=10000)  # Limit rows
            else:
                df = pd.read_excel(file_path, nrows=10000)
            
            text_parts = []
            
            # Add column description
            columns = df.columns.tolist()
            text_parts.append(f"Columns: {', '.join(str(c) for c in columns)}")
            
            # Add data summary
            text_parts.append(f"Total rows: {len(df)}")
            
            # Convert to text representation
            for idx, row in df.head(500).iterrows():  # Limit to 500 rows for text
                row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                text_parts.append(row_text)
            
            # Add statistical summary for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("\nStatistical Summary:")
                for col in numeric_cols:
                    stats = df[col].describe()
                    text_parts.append(f"{col}: mean={stats['mean']:.2f}, min={stats['min']:.2f}, max={stats['max']:.2f}")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting tabular data: {e}")
            raise
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise
    
    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        return self._extract_text(file_path)
    
    def _extract_json(self, file_path: str) -> str:
        """Extract text from JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to readable text
            return self._json_to_text(data)
            
        except Exception as e:
            logger.error(f"Error extracting JSON: {e}")
            raise
    
    def _json_to_text(self, data: Any, prefix: str = "") -> str:
        """Convert JSON data to readable text."""
        if isinstance(data, dict):
            parts = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    parts.append(f"{prefix}{key}:")
                    parts.append(self._json_to_text(value, prefix + "  "))
                else:
                    parts.append(f"{prefix}{key}: {value}")
            return "\n".join(parts)
        elif isinstance(data, list):
            parts = []
            for i, item in enumerate(data[:100]):  # Limit to 100 items
                if isinstance(item, (dict, list)):
                    parts.append(f"{prefix}Item {i + 1}:")
                    parts.append(self._json_to_text(item, prefix + "  "))
                else:
                    parts.append(f"{prefix}- {item}")
            return "\n".join(parts)
        else:
            return str(data)
    
    def _preprocess_text(self, text: str) -> str:
        """
        Preprocess text with energy domain optimization.
        - Normalize units (MW, GW, kW)
        - Standardize date formats
        - Clean up whitespace
        """
        # Normalize power units
        text = re.sub(r'(\d+)\s*MW', r'\1 MW', text, flags=re.IGNORECASE)
        text = re.sub(r'(\d+)\s*GW', r'\1 GW', text, flags=re.IGNORECASE)
        text = re.sub(r'(\d+)\s*kW', r'\1 kW', text, flags=re.IGNORECASE)
        text = re.sub(r'(\d+)\s*MWh', r'\1 MWh', text, flags=re.IGNORECASE)
        
        # Normalize voltage units
        text = re.sub(r'(\d+)\s*kV', r'\1 kV', text, flags=re.IGNORECASE)
        text = re.sub(r'(\d+)\s*V\b', r'\1 V', text)
        
        # Normalize frequency
        text = re.sub(r'(\d+)\s*Hz', r'\1 Hz', text, flags=re.IGNORECASE)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def _detect_domain(self, text: str) -> Optional[str]:
        """Detect energy domain from text content."""
        text_lower = text.lower()
        
        domain_scores = {}
        for domain, keywords in ENERGY_DOMAINS.items():
            score = sum(1 for kw in keywords if kw in text_lower)
            if score > 0:
                domain_scores[domain] = score
        
        if domain_scores:
            return max(domain_scores, key=domain_scores.get)
        return None
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        if tokenizer:
            return len(tokenizer.encode(text))
        else:
            # Approximate: ~4 characters per token
            return len(text) // 4
    
    def _chunk_text(
        self,
        text: str,
        doc_id: str,
        filename: str,
        source: Optional[str],
        security_level: SecurityLevel,
        domain: Optional[str],
        file_type: str
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        chunks = []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_tokens = 0
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_tokens = self._count_tokens(para)
            
            # If single paragraph exceeds chunk size, split it
            if para_tokens > self.chunk_size:
                # Save current chunk if exists
                if current_chunk and current_tokens >= self.min_chunk_size:
                    chunks.append(self._create_chunk(
                        text=current_chunk,
                        doc_id=doc_id,
                        filename=filename,
                        source=source,
                        chunk_index=chunk_index,
                        security_level=security_level,
                        domain=domain,
                        file_type=file_type
                    ))
                    chunk_index += 1
                    current_chunk = ""
                    current_tokens = 0
                
                # Split large paragraph by sentences
                sentences = re.split(r'(?<=[.!?])\s+', para)
                for sentence in sentences:
                    sentence_tokens = self._count_tokens(sentence)
                    
                    if current_tokens + sentence_tokens <= self.chunk_size:
                        current_chunk = (current_chunk + " " + sentence).strip()
                        current_tokens += sentence_tokens
                    else:
                        if current_chunk and current_tokens >= self.min_chunk_size:
                            chunks.append(self._create_chunk(
                                text=current_chunk,
                                doc_id=doc_id,
                                filename=filename,
                                source=source,
                                chunk_index=chunk_index,
                                security_level=security_level,
                                domain=domain,
                                file_type=file_type
                            ))
                            chunk_index += 1
                        
                        # Start new chunk with overlap
                        if chunks:
                            overlap_text = self._get_overlap(current_chunk)
                            current_chunk = (overlap_text + " " + sentence).strip()
                        else:
                            current_chunk = sentence
                        current_tokens = self._count_tokens(current_chunk)
            
            elif current_tokens + para_tokens <= self.chunk_size:
                current_chunk = (current_chunk + "\n\n" + para).strip()
                current_tokens += para_tokens
            else:
                # Save current chunk
                if current_chunk and current_tokens >= self.min_chunk_size:
                    chunks.append(self._create_chunk(
                        text=current_chunk,
                        doc_id=doc_id,
                        filename=filename,
                        source=source,
                        chunk_index=chunk_index,
                        security_level=security_level,
                        domain=domain,
                        file_type=file_type
                    ))
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap(current_chunk)
                current_chunk = (overlap_text + "\n\n" + para).strip()
                current_tokens = self._count_tokens(current_chunk)
        
        # Don't forget the last chunk
        if current_chunk and current_tokens >= self.min_chunk_size:
            chunks.append(self._create_chunk(
                text=current_chunk,
                doc_id=doc_id,
                filename=filename,
                source=source,
                chunk_index=chunk_index,
                security_level=security_level,
                domain=domain,
                file_type=file_type
            ))
        
        # Update total_chunks in metadata
        total_chunks = len(chunks)
        for chunk in chunks:
            chunk.metadata.total_chunks = total_chunks
        
        return chunks
    
    def _get_overlap(self, text: str) -> str:
        """Get overlap text from the end of a chunk."""
        if not text:
            return ""
        
        if tokenizer:
            tokens = tokenizer.encode(text)
            overlap_tokens = tokens[-self.chunk_overlap:]
            return tokenizer.decode(overlap_tokens)
        else:
            # Approximate overlap by characters
            chars = self.chunk_overlap * 4
            return text[-chars:] if len(text) > chars else text
    
    def _create_chunk(
        self,
        text: str,
        doc_id: str,
        filename: str,
        source: Optional[str],
        chunk_index: int,
        security_level: SecurityLevel,
        domain: Optional[str],
        file_type: str
    ) -> DocumentChunk:
        """Create a DocumentChunk object."""
        # Generate a proper UUID for Qdrant compatibility
        chunk_id = str(uuid.uuid4())
        
        metadata = ChunkMetadata(
            doc_id=doc_id,
            filename=filename,
            source=source,
            chunk_index=chunk_index,
            total_chunks=0,  # Updated later
            security_level=security_level,
            created_at=datetime.now(),
            domain=domain,
            file_type=file_type
        )
        
        return DocumentChunk(
            id=chunk_id,
            text=text,
            metadata=metadata
        )


# ===== Global Processor Instance =====

document_processor = DocumentProcessor()


def get_supported_extensions() -> List[str]:
    """Get list of supported file extensions."""
    return list(SUPPORTED_EXTENSIONS.keys())
